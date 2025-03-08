from operator import index

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import re
import json
import time
from io import StringIO, BytesIO
import requests
from PIL import Image
from docx import Document
from docx.shared import Inches
import base64
# New imports for audio recording and processing
import sounddevice as sd
import soundfile as sf
import numpy as np
import tempfile
from openai import OpenAI
import guardrails as gd
from guardrails import Guard, OnFailAction
from guardrails.validator_base import Validator, PassResult, FailResult, register_validator
from fuzzywuzzy import fuzz

# Streamlit app configuration
st.set_page_config(
    page_title="Perrigo GenAI Answer Bot",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

logo = Image.open("Images/perrigo-logo.png")
st.image(logo, width=120)

# Custom CSS
st.markdown("""
    <style>
    .stAlert {
        padding: 1rem;
        margin-bottom: 1rem;
    }
    .st-emotion-cache-16idsys p {
        font-size: 1.1rem;
    }
    </style>
    """, unsafe_allow_html=True)

import parameter_values
import cost_cosnsolidation
import static


# Initialize session state for audio recording and transcribed text
if 'recording' not in st.session_state:
    st.session_state.recording = False
if 'audio_data' not in st.session_state:
    st.session_state.audio_data = None
if 'transcribed_text' not in st.session_state:
    st.session_state.transcribed_text = ""

if 'messages' not in st.session_state:
    st.session_state.messages = []


# Function to record audio
def record_audio(duration=8, sample_rate=16000):
    """Record audio for a specified duration."""
    audio_data = sd.rec(int(duration * sample_rate), samplerate=sample_rate, channels=1)
    sd.wait()
    return audio_data


# Function to transcribe audio using Whisper
def transcribe_audio(audio_file_path, api_key):
    """Transcribe audio file using OpenAI Whisper API."""
    try:
        client = OpenAI(api_key=api_key)
        with open(audio_file_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                language="en"
            )
        return transcript.text
    except Exception as e:
        st.error(f"Error in transcription: {str(e)}")
        return None


# Initialize session state
if 'chat_history' not in st.session_state:
    st.session_state.chat_history = []
if 'initialized' not in st.session_state:
    st.session_state.initialized = False
if 'current_data_source' not in st.session_state:
    st.session_state.current_data_source = None


def reset_app_state():
    """Reset the app state when data source changes"""
    st.session_state.initialized = False
    if 'df' in st.session_state:
        del st.session_state.df


def save_figure_to_image(fig):
    """Convert matplotlib figure to bytes for Word document."""
    buf = BytesIO()
    fig.savefig(buf, format='png', bbox_inches='tight')
    buf.seek(0)
    return buf


def create_word_document(chat_history):
    """Create a Word document from the analysis history."""
    doc = Document()
    doc.add_heading('Data Analysis History Report', 0)

    # Add generation timestamp
    doc.add_paragraph(f'Generated on: {time.strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph('-' * 50)

    # Add each analysis to the document
    for idx, chat in enumerate(reversed(chat_history), 1):
        # Add query section with data source
        doc.add_heading(f'Query: {chat["query"]}', level=1)
        doc.add_paragraph(f'Data Source: {chat.get("data_source", "Not specified")}')

        # Rest of the function remains the same

        if chat['approach']:
            doc.add_heading('Approach:', level=2)
            doc.add_paragraph(chat['approach'])

        if chat['answer']:
            doc.add_heading('Results:', level=2)
            doc.add_paragraph(chat['answer'])

        if chat['figure']:
            doc.add_heading('Visualization:', level=2)
            image_stream = save_figure_to_image(chat['figure'])
            doc.add_picture(image_stream, width=Inches(6))

        doc.add_paragraph('-' * 50)

    return doc


def download_word_doc():
    """Create and return a download link for the Word document."""
    if not st.session_state.chat_history:
        st.warning("No analysis history to export!")
        return

    # Create Word document
    doc = create_word_document(st.session_state.chat_history)

    # Save document to bytes
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # Create download button
    st.download_button(
        label="📥 Download Analysis History",
        data=doc_bytes.getvalue(),
        file_name=f"analysis_history_{time.strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def extract_code_segments(response_text):
    """Extract code segments from the API response using regex."""
    segments = {}

    # Extract approach section
    approach_match = re.search(r'<approach>(.*?)</approach>', response_text, re.DOTALL)
    if approach_match:
        segments['approach'] = approach_match.group(1).strip()

    # Extract content between <code> tags
    code_match = re.search(r'<code>(.*?)</code>', response_text, re.DOTALL)
    if code_match:
        segments['code'] = code_match.group(1).strip()

    # Extract content between <chart> tags
    chart_match = re.search(r'<chart>(.*?)</chart>', response_text, re.DOTALL)
    if chart_match:
        segments['chart'] = chart_match.group(1).strip()

    # Extract content between <answer> tags
    answer_match = re.search(r'<answer>(.*?)</answer>', response_text, re.DOTALL)
    if answer_match:
        segments['answer'] = answer_match.group(1).strip()

    return segments


def execute_analysis(df, response_text):
    """Execute the extracted code segments on the provided dataframe and store formatted answer."""
    results = {
        'approach': None,
        'answer': None,
        'figure': None,
        'code': None,
        'chart_code': None
    }

    try:
        # Extract code segments
        segments = extract_code_segments(response_text)

        if not segments:
            st.error("No code segments found in the response")
            return results

        # Store the approach and raw code
        if 'approach' in segments:
            results['approach'] = segments['approach']
        if 'code' in segments:
            results['code'] = segments['code']
        if 'chart' in segments:
            results['chart_code'] = segments['chart']

        # Create a single namespace for all executions
        namespace = {'df': df, 'pd': pd, 'plt': plt, 'sns': sns}

        # Execute analysis code and answer template
        if 'code' in segments and 'answer' in segments:
            # Properly dedent the code before execution
            code_lines = segments['code'].strip().split('\n')
            # Find minimum indentation
            min_indent = float('inf')
            for line in code_lines:
                if line.strip():  # Skip empty lines
                    indent = len(line) - len(line.lstrip())
                    min_indent = min(min_indent, indent)
            # Remove consistent indentation
            dedented_code = '\n'.join(line[min_indent:] if line.strip() else ''
                                      for line in code_lines)

            # Combine code with answer template
            combined_code = f"""
{dedented_code}

# Format the answer template
answer_text = f'''{segments['answer']}'''
"""
            exec(combined_code, namespace)
            results['answer'] = namespace.get('answer_text')

        # Execute chart code if present
        if 'chart' in segments:
            # Properly dedent the chart code
            chart_lines = segments['chart'].strip().split('\n')
            # Find minimum indentation
            min_indent = float('inf')
            for line in chart_lines:
                if line.strip():  # Skip empty lines
                    indent = len(line) - len(line.lstrip())
                    min_indent = min(min_indent, indent)
            # Remove consistent indentation
            dedented_chart = '\n'.join(line[min_indent:] if line.strip() else ''
                                       for line in chart_lines)

            plt.figure(figsize=(10, 6))
            exec(dedented_chart, namespace)
            fig = plt.gcf()
            results['figure'] = fig
            plt.close()

        return results

    except Exception as e:
        st.error(f"Error during execution: {str(e)}")
        return results


def get_prompt_file(data_source):
    """Return the appropriate prompt file based on the data source."""
    prompt_mapping = {
        'Outbound_Data.csv': 'Prompts/Prompt1.txt',
        'Inventory_Batch.csv': 'Prompts/Prompt2.txt',
        'Inbound_Data.csv': 'Prompts/Prompt3.txt'
    }
    return prompt_mapping.get(data_source)


def finding_intent(query, api_key):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }
    payload = {
        "model": "gpt-4o",
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"""
You are an AI assistant that identifies the intent behind a user's query. The intent can be either data analysis or cost optimization. Use the following guidelines to classify the intent:
Data Analysis:
Queries focus on analyzing datasets, identifying trends, distributions, patterns, or relationships.
Examples:
"What is the monthly trend in total cost?"
"Identify the distribution of cost per pallet, is it normally distributed?"
"What is the average shelf life remaining for each product category?"
"Generate a heat map of inventory levels across different storage locations."
Cost Optimization:
Queries focus on strategies, scenarios, or calculations to minimize costs or improve efficiency.
Examples:
"What are the cost savings for user X if we consolidate shipments?"
"How can I optimize shipment costs for user X?"
"What is the best consolidation window for user X to minimize costs?"
"What happens if we increase truck capacity by 10%?"
Here is the user's query:
<query> {query} </query>
Task:
Classify the query intent as either data analysis or cost optimization and provide a brief reason for your classification and include this inside the <intent> tags.
Format:
<intent> [Data Analysis / Cost Optimization] </intent>
Example Outputs:
Query:
"What is the monthly trend in total cost?"
<intent>Data Analysis</intent>
Query:
"What are the cost savings for user X if we consolidate shipments?"
<intent>Cost Optimization</intent>

                    """
                    }
                ]
            }
        ],
        "max_tokens": 4096,
        "temperature": 0
    }
    try:
        response = requests.post("https://api.openai.com/v1/chat/completions",
                                 headers=headers,
                                 json=payload)

        if response.status_code != 200:
            st.error(f"Error: Received status code {response.status_code}")
            st.error(f"Response content: {response.text}")
            return None

        response_json = response.json()
        response_content = response_json['choices'][0]['message']['content']

        # Extract intent section
        intent_match = re.search(r'<intent>(.*?)</intent>', response_content, re.DOTALL)
        if intent_match:
            intent = intent_match.group(1).strip()

        return intent

    except Exception as e:
        st.error(f"Error during analysis: {e}")
        return None

Description_of_Data = """
This dataset contains information about shipping orders, including:
- Order ID
- Number of pallets
- Shipping cost
- Distance traveled
- Product type (PROD_TYPE)
- Destination postcode (SHORT_POSTCODE)
- Cost savings

The dataset can be used to answer questions about:
- Total orders and pallets
- Average shipping cost and distance
- Most common product types
- Delivery performance over time
- Cost savings for customers or postcodes
"""

sample_queries = [
    "What is the monthly trend in total cost?",
    "What is the average cost per pallet for each PROD TYPE and how does it vary across the following SHORT_POSTCODE regions: CV, NG, NN, RG?",
    "Which customer has the highest total shipping cost over time, and how does its cost trend vary by month?",
    "Identify the SHORT_POSTCODE areas with the highest total shipping costs and also mention their cost per pallet.",
    "For ambient product type, which are the top 5 customers with total orders > 10 and highest standard deviation in cost per pallet?",
    "How does the cost per order vary with distance within each PROD TYPE?",
    "How can I optimize the shipment costs for user ALLOGA UK in January?",
    "What is the cost saving for postal codes NG, LU and NN in Feb 2024 if we consolidate?",
    "How can we save spends for Tesco grocery dept in the month of March?",
    "How can I improve my cost savings for first quarter of 2024 for customer Sainsbury's Supermarket?",
    "How can we optimise spends for customer Boots Company PLC for the month of Feb to maximise efficiency?",
    "What are the best ways to reduce shipment costs for B & M Ltd Retail while maintaining service efficiency?"
]


# @register_validator(name="check_relevance", data_type="string")
# class RelevanceValidator(Validator):
#     def __init__(self, on_fail=None):
#         super().__init__(on_fail=on_fail)
#         self.dataset_description = Description_of_Data.lower()
#         self.sample_questions = [q.lower() for q in sample_queries]
#
#         self.negative_keywords = [
#             "advertisement", "google", "social media", "marketing",
#             "website", "seo", "click", "impression", "campaign"
#         ]
#
#         # 🟢 Required context keywords from shipping domain
#         self.required_context = [
#             "shipping", "pallet", "postcode", "prod type",
#             "customer", "consolidate", "distance", "destination",
#             "cost saving", "order id", "shipment"
#         ]
#
#     def _validate(self, value, metadata=None):
#         query = value.lower()
#
#         if any(nk in query for nk in self.negative_keywords):
#             #st.write(f"🚫 Blocked by negative keywords")
#             return FailResult("Query about unrelated domain (ads/marketing)")
#
#             # 2️⃣ Second check: Must contain shipping-related context
#         if not any(rc in query for rc in self.required_context):
#             #st.write(f"🚫 Missing required shipping context")
#             return FailResult("Query lacks shipping-specific terminology")
#
#             # 3️⃣ Third check: Dataset keyword alignment
#         dataset_keywords = ["cost saving", "order", "distance",
#                             "prod_type", "short_postcode", "customer" ,"postcodes"]
#         if not any(kw in query for kw in dataset_keywords):
#             #st.write(f"🚫 No dataset-specific keywords found")
#             return FailResult("Query doesn't reference dataset columns")
#
#         max_similarity = max(fuzz.ratio(query, q) for q in self.sample_questions)
#         if max_similarity < 60:  # More strict threshold
#             st.write(f"🚫 No similar sample questions ({max_similarity}%)")
#             return FailResult("Query pattern doesn't match known use cases")
#
#         keywords = self.dataset_description.replace("\n", " ").split(" ")
#         keywords = [word.strip(".,-") for word in keywords if word.strip(".,-")]
#
#         # Check if any keyword is in the query
#         if any(keyword in query for keyword in keywords):
#             st.write("✅ Query matches dataset description.")
#             return PassResult()
#
#         # ✅ Fuzzy Matching with Sample Questions (Threshold 80 for High Similarity)
#         for question in self.sample_questions:
#             similarity_score = fuzz.ratio(query, question)
#             if similarity_score > 80:  # 80% or higher similarity
#                 st.write(f"✅ Query matches a sample question ({similarity_score}% similarity).")
#                 return PassResult()
#
#         st.write("🚫 Query is NOT relevant.")  # Debug log
#         return FailResult(error_message="Query is unrelated to your dataset.")

@register_validator(name="check_relevance", data_type="string")
class RelevanceValidator(Validator):
    def __init__(self, on_fail=None):
        super().__init__(on_fail=on_fail)
        self.dataset_description = Description_of_Data.lower()
        self.sample_questions = [q.lower() for q in sample_queries]

        # Negative keywords to block irrelevant questions
        self.negative_keywords = [
            "advertisement", "google", "social media", "marketing",
            "website", "click", "impression", "campaign"
        ]

        # Required context keywords from shipping domain
        self.required_context = [
            "shipping", "pallet", "postcode", "prod type",
            "customer", "consolidate", "distance", "destination",
            "cost saving", "order id", "shipment"
        ]

        # Dataset-specific keywords
        self.dataset_keywords = [
            "cost saving", "order", "distance",
            "prod_type", "short_postcode", "customer", "postcodes"
        ]

    def _validate(self, value, metadata=None):
        query = value.lower()

        # 1️⃣ First check: Block questions with negative keywords
        if any(nk in query for nk in self.negative_keywords):
            st.write(f"🚫 Blocked by negative keywords")
            return FailResult("Query about unrelated domain (ads/marketing)")

        # 2️⃣ Second check: Must contain at least TWO shipping-related keywords
        shipping_keywords_found = sum(rc in query for rc in self.required_context)
        if shipping_keywords_found < 2:  # Require at least 2 relevant keywords
            st.write(f"🚫 Missing required shipping context")
            return FailResult("Query lacks sufficient shipping-specific terminology")

        # 3️⃣ Third check: Dataset keyword alignment
        dataset_keywords_found = sum(kw in query for kw in self.dataset_keywords)
        if dataset_keywords_found < 2:  # Require at least 2 dataset-specific keywords
            st.write(f"🚫 No dataset-specific keywords found")
            return FailResult("Query doesn't reference dataset columns")

        # 4️⃣ Fourth check: Fuzzy Matching with Sample Questions (Strict Threshold)
        max_similarity = max(fuzz.ratio(query, q) for q in self.sample_questions)
        if max_similarity < 70:  # Increased threshold for stricter matching
            st.write(f"🚫 No similar sample questions ({max_similarity}%)")
            return FailResult("Query pattern doesn't match known use cases")

        # 5️⃣ Fifth check: Semantic relevance to dataset description
        keywords = self.dataset_description.replace("\n", " ").split(" ")
        keywords = [word.strip(".,-") for word in keywords if word.strip(".,-")]
        if not any(keyword in query for keyword in keywords):
            st.write("🚫 Query doesn't match dataset description.")
            return FailResult("Query is unrelated to your dataset.")

        # ✅ If all checks pass, the query is relevant
        st.write("✅ Query is relevant.")
        return PassResult()


def analyze_data_with_execution(df, question, api_key, data_source):
    # Get the appropriate prompt file based on data source
    validators = [RelevanceValidator(on_fail=OnFailAction.EXCEPTION)]

    guard = Guard.for_string(validators)
    #guard = Guard.from_string(validators=[RelevanceValidator(on_fail=OnFailAction.EXCEPTION)])

    # Validate the relevance of the question
    try:
        guard.validate(question, metadata={})
    except Exception as e:
        st.write("🚫 The question is NOT relevant to the dataset. Please ask a relevant question.")
        return None  # Stop execution if the question is not relevant

    # If the question is relevant, proceed with the analysis

    prompt_file = get_prompt_file(data_source)

    if not prompt_file:
        st.error("Unable to determine prompt file for the selected data source!")
        return None

    # Read the prompt template from file
    try:
        with open(prompt_file, 'r') as file:
            data_description = file.read().strip()
    except FileNotFoundError:
        st.error(f"{prompt_file} file not found!")
        return None
    except Exception as e:
        st.error(f"Error reading {prompt_file}: {str(e)}")
        return None

    # Append the user's question to session messages
    st.session_state.messages.append({"role": "user", "content": question})

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}"
    }

    payload = {
        "model": "gpt-4o",
        "messages": [
            *st.session_state.messages,
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": f"""
You are an AI assistant tasked with analyzing a dataset to provide code for calculating the final answer and generating relevant visualization.
I will provide you with the data in dataframe format, as well as a question to answer based on the data.

{data_description}

Here is the question I would like you to answer using this data:
<question>
{question}
</question>

To answer this, first think through your approach inside <approach> tags. Break down the steps you
will need to take and consider which columns of the data will be most relevant. Here is an example:
<approach>
To answer this question, I will need to:
1. Calculate the total number of orders and pallets across all rows
2. Determine the average distance and cost per order
3. Identify the most common PROD_TYPE and SHORT_POSTCODE
</approach>

Then, write the Python code needed to analyze the data and calculate the final answer inside <code> tags. Assume input dataframe as 'df'
Be sure to include any necessary data manipulation, aggregations, filtering, etc. Return only the Python code without any explanation or markdown formatting.
For decimal answers round them to 1 decimal place.

Generate Python code using matplotlib and/or seaborn to create an appropriate chart to visualize the relevant data and support your answer.
For example if user is asking for postcode with highest cost then a relevant chart can be a bar chart showing top 10 postcodes with highest total cost arranged in decreasing order.
Specify the chart code inside <chart> tags.
When working with dates:

Always convert dates to datetime using pd.to_datetime() with explicit format
For grouping by month, use dt.strftime('%Y-%m') instead of dt.to_period()
Sort date-based results chronologically before plotting

The visualization code should follow these guidelines:

Start with these required imports:

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns

Use standard chart setup:
# Set figure size and style
plt.figure(figsize=(8, 5))
# Set seaborn default style and color palette
sns.set_theme(style="whitegrid")  
sns.set_palette('pastel')

For time-based charts:


Use string dates on x-axis (converted using strftime)
Rotate labels: plt.xticks(rotation=45, ha='right')
Add gridlines: plt.grid(True, alpha=0.3)

For large numbers:
Format y-axis with K/M suffixes using:

Always include:

Clear title (plt.title())
Axis labels (plt.xlabel(), plt.ylabel())
plt.tight_layout() at the end


For specific chart types:

Time series: sns.lineplot() with marker='o'
Rankings: sns.barplot() with descending sort
Comparisons: sns.barplot() or sns.boxplot()
Distributions: sns.histplot() or sns.kdeplot()

Return only the Python code without any explanation or markdown formatting.

Finally, provide the answer to the question in natural language inside <answer> tags. Be sure to
include any key variables that you calculated in the code inside {{}}.

                    """
                    }
                ]
            }
        ],
        "max_tokens": 4096,
        "temperature": 0
    }

    try:
        response = requests.post("https://api.openai.com/v1/chat/completions",
                                 headers=headers,
                                 json=payload)

        if response.status_code != 200:
            st.error(f"Error: Received status code {response.status_code}")
            st.error(f"Response content: {response.text}")
            return None

        response_json = response.json()
        response_content = response_json['choices'][0]['message']['content']

            # Save the assistant's response to the session state
        st.session_state.messages.append({"role": "assistant", "content": response_content})
        # Execute the code segments and get results
        results = execute_analysis(df, response_content)

        return results

    except Exception as e:
        st.error(f"Error during analysis: {e}")
        return None


def load_data_file(filename):
    """Load a CSV data file with automatic parsing of date columns."""
    try:
        # Load data without parsing dates first
        data = pd.read_csv(filename)

        # Identify columns with "date" in their name and parse them as dates
        date_columns = [col for col in data.columns if 'date' in col.lower()]

        # Reload data with date parsing for identified columns
        return pd.read_csv(filename, parse_dates=date_columns, dayfirst=True)

    except Exception as e:
        st.error(f"Error loading {filename}: {str(e)}")
        return None


def display_analysis_results(results):
    """Display the analysis results in a structured format."""

    if results['approach']:
        st.subheader("Approach")
        st.write(results['approach'])

    if results['answer']:
        st.subheader("Analysis Results")
        st.write(results['answer'])

    if results['figure']:
        st.subheader("Visualization")
        st.pyplot(results['figure'])

    # Display code segments in expandable sections
    if results['code'] or results['chart_code']:
        st.subheader("Code Segments")

        if results['code']:
            with st.expander("Show Analysis Code"):
                st.code(results['code'], language='python')

        if results['chart_code']:
            with st.expander("Show Visualization Code"):
                st.code(results['chart_code'], language='python')


def get_sample_queries(data_source):
    """Return appropriate sample queries based on the selected data source."""
    queries = {
        'Outbound_Data.csv': [
            "What is the monthly trend in total cost?",
            "What is the average cost per pallet for each PROD TYPE and how does it vary across the following SHORT_POSTCODE regions: CV, NG, NN, RG?",
            "Which customer has the highest total shipping cost over time, and how does its cost trend vary by month?",
            "Identify the SHORT_POSTCODE areas with the highest total shipping costs and also mention their cost per pallet.",
            "For ambient product type, which are the top 5 customers with total orders > 10 and highest standard deviation in cost per pallet?",
            "How does the cost per order vary with distance within each PROD TYPE?",
            "How can I optimize the shipment costs for user ALLOGA UK in January?",
            "What is the cost saving for postal codes NG, LU and NN in Feb 2024 if we consolidate?",
            "How can we save spends for Tesco grocery dept in the month of March?",
            "How can I improve my cost savings for first quarter of 2024 for customer Sainsbury's Supermarket?",
            "How can we optimise spends for customer Boots Company PLC for the month of Feb to maximise efficiency?",
            "What are the best ways to reduce shipment costs for B & M Ltd Retail while maintaining service efficiency?"
         ],
        'Inventory_Batch.csv': [
            "What is the total inventory value by product category?",
            "Which products have inventory levels below their safety stock?",
            "What is the monthly trend in inventory turnover rate?",
            "Show the age distribution of current inventory batches",
            "Which are the top 10 products by storage cost?",
            "What is the average shelf life remaining for each product category?",
            "Identify products with excess inventory (more than 120% of max stock level)",
            "What is the weekly trend in inventory receipts vs. withdrawals?",
            "Generate a heat map of inventory levels across different storage locations",
            "Which products have the highest holding costs in the last quarter?",
            "Show the distribution of batch sizes by product category",
            "What is the correlation between product value and storage duration?",
            "Identify seasonal patterns in inventory levels for the top 5 products",
            "Calculate and visualize the inventory accuracy rates by location",
            "What is the average time between receipt and first withdrawal for each product?",
            "Show the distribution of inventory value across different temperature zones",
            "Which products have the highest stock rotation frequency?",
            "Generate a Pareto chart of inventory value by product category",
            "What is the trend in average days of inventory on hand?"
        ],
        'Inbound_Data.csv': [
            "What is the utilization in each tradelane for top 15 tradelane by pallets?",
            "What is the total cost in each route from Nov 2023 to Jan 2024? Consider top 10 routes with highest total pallets.",
            "What is the monthly trend of above metrics?",
            "What is the cost breakdown by Company?",
            "What is the proportion of FTL/LTL by route?",
            "What is the Pallet per Order?",
            "What is the cost per pallet?",
            "What is the cost per order?",
            "What is the average lead time by tradelane/tradeline/route?",
            "Which routes/delivery supplier/delivery groups charge higher fuel costs?",
            "Which routes/delivery supplier/delivery groups have higher % of late delivery?",
            "Which routes/delivery supplier/delivery groups have higher % of late collection?",
            "What is the average delay in delivery on a particular route by delivery supplier?",
            "What is the average delay in collection on a particular route by delivery supplier?"
        ]
    }
    return queries.get(data_source, [])

# Initialize session state
if "show_cost_options" not in st.session_state:
    st.session_state.show_cost_options = False
if "cost_consolidation_type" not in st.session_state:
    st.session_state.cost_consolidation_type = None
if "results" not in st.session_state:
    st.session_state.results = None
if "dynamic_clicked" not in st.session_state:
    st.session_state.dynamic_clicked = False
if "static_clicked" not in st.session_state:
    st.session_state.static_clicked = False
if "total_shipment_capacity" not in st.session_state:
    st.session_state.total_shipment_capacity = 46
if "shipment_window_range" not in st.session_state:
    st.session_state.shipment_window_range = (2, 7)
if 'delivery_options' not in st.session_state:
    st.session_state.delivery_options = None
if 'selected_delivery_option' not in st.session_state:
    st.session_state.selected_delivery_option = None

def reset_session_states():
    st.session_state.show_cost_options = False
    st.session_state.cost_consolidation_type = None
    st.session_state.dynamic_clicked = False
    st.session_state.results = None


def main():
    st.title("GenAI Answer Bot")

    # Define available data files
    data_files = {
        'Outbound_Data.csv': 'Data/Outbound_Data.csv',
        'Inventory_Batch.csv': 'Data/Inventory_Batch.csv',
        'Inbound_Data.csv': 'Data/Inbound_Data.csv'
    }

    # Sidebar configuration
    with st.sidebar:
        st.header("⚙️ Configuration")

        # API Key input
        st.subheader("1. API Key")
        api_key = st.text_input("Enter OpenAI API Key:", type="password")

        # Data source selection
        st.subheader("2. Data Source")
        data_source = st.radio(
            "Choose Data Source:",
            # list(data_files.keys()) + ["Upload Custom File"],
            list(data_files.keys()),
            disabled=False,
            index=0
        )

        # Reset state if data source changes
        if st.session_state.current_data_source != data_source:
            st.session_state.current_data_source = data_source
            reset_app_state()

        df = None
        if data_source in data_files:
            df = load_data_file(data_files[data_source])
            if df is not None:
                st.success(f"{data_source} loaded successfully!")
                st.session_state.df = df
        else:
            uploaded_file = st.file_uploader("Upload CSV file", type=['csv'])
            if uploaded_file:
                try:
                    temp_df = pd.read_csv(uploaded_file)

                    # Identify columns with "date" in their name and parse those as dates
                    date_columns = [col for col in temp_df.columns if 'date' in col.lower()]

                    # Reload the data with date parsing for identified columns
                    df = pd.read_csv(uploaded_file, parse_dates=date_columns, dayfirst=True)

                    st.success("Custom file loaded successfully!")
                    st.session_state.df = df
                except Exception as e:
                    st.error(f"Error loading custom file: {str(e)}")

    # Main content area
    if not api_key:
        st.info("Please enter your OpenAI API key in the sidebar to get started.")
        return

    if 'df' not in st.session_state:
        if data_source in data_files:
            st.error(f"Data file not found. Please check if '{data_source}' exists.")
        else:
            st.info("Please upload your CSV file in the sidebar.")
        return

    # Display sample data
    with st.expander("📊 View Sample Data"):
        display_df = st.session_state.df.copy()

        # Identify and format all datetime columns
        date_columns = display_df.select_dtypes(include=['datetime64']).columns
        for date_col in date_columns:
            display_df[date_col] = display_df[date_col].dt.strftime('%d-%m-%Y')

        display_df = display_df.set_index(display_df.columns[0])
        st.dataframe(display_df.head(), use_container_width=True)

    # Query interface
    st.subheader("💬 Ask Questions About Your Data")

    # Get sample queries based on selected data source
    sample_queries = get_sample_queries(st.session_state.current_data_source)

    selected_query = st.selectbox(
        "Select a sample query or write your own below:",
        [""] + sample_queries,
        key="query_select"
    )

    # Create columns for text input and microphone button
    col1, col2 = st.columns([18, 1])

    # Use the transcribed text from session state if available, otherwise use selected query
    initial_value = st.session_state.transcribed_text if st.session_state.transcribed_text else selected_query

    with col1:
        query = st.text_area(
            "Enter your query:",
            value=initial_value,
            height=100,
            key="query_input"
        )

    # Create a placeholder for error messages before the microphone button
    error_placeholder = st.empty()

    with col2:
        st.write("")
        st.write("")
        # Add microphone button with progress bar
        if st.button("🎙️", help="Start Recording"):
            try:
                # Create a progress bar
                progress_bar = st.progress(0)
                status_text = st.empty()

                # Record audio while updating progress
                audio_data = None
                sample_rate = 16000
                duration = 8

                # Start recording
                audio_data = sd.rec(int(duration * sample_rate), samplerate=sample_rate, channels=1)

                # Update progress bar while recording
                for i in range(duration):
                    # Update progress
                    progress = (i + 1) / duration
                    progress_bar.progress(progress)
                    time.sleep(1)

                sd.wait()  # Wait for recording to complete

                # Clear progress bar and status
                progress_bar.empty()
                status_text.empty()

                st.session_state.audio_data = audio_data

                # Save audio to temporary file
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
                    sf.write(temp_audio.name, audio_data, 16000)

                    # Transcribe audio
                    transcribed_text = transcribe_audio(temp_audio.name, api_key)

                    if transcribed_text:
                        # Store transcribed text in session state
                        st.session_state.transcribed_text = transcribed_text
                        st.rerun()
                    else:
                        # Use the placeholder to display the error
                        error_placeholder.error("Failed to transcribe audio. Please try again.")

            except Exception as e:
                # Use the placeholder to display the error
                error_placeholder.error(f"Error recording audio: {str(e)}")


    col1, col2 = st.columns([1, 5])
    with col1:
        submit_button = st.button("🔍 Analyze")

    if submit_button and query:

        # Move time tracking and spinner to encompass both analysis and display
        with st.spinner("Analyzing data "):

            intent = finding_intent(query=query, api_key=api_key)
            if "last_intent" not in st.session_state:
                st.session_state.last_intent = intent

            # If the intent has changed, reset the session states
            if intent != st.session_state.last_intent:
                reset_session_states()
                st.session_state.last_intent = intent

            if intent == 'Data Analysis':
                # Perform analysis
                #st.write("It is Data Analysis prompt\n")
                results = analyze_data_with_execution(
                    st.session_state.df,
                    query,
                    api_key,
                    st.session_state.current_data_source
                )
            elif intent == 'Cost Optimization':
                results = None
                if 'results' in st.session_state:
                    del st.session_state.results
                if 'selected_postcodes' in st.session_state:
                    del st.session_state.selected_postcodes
                if 'selected_customers' in st.session_state:
                    del st.session_state.selected_customers
                if 'parameters' in st.session_state:
                    del st.session_state.parameters

                st.session_state.show_cost_options = True

            if results and intent != 'Cost Optimization':
                # Display results inside the spinner context
                display_analysis_results(results)

                # Store in chat history
                chat_entry = {
                    "data_source": st.session_state.current_data_source,
                    "query": query,
                    "approach": results['approach'],
                    "answer": results['answer'],
                    "figure": results['figure'],
                    "code": results['code'],
                    "chart_code": results['chart_code'],
                }

                if not st.session_state.chat_history or st.session_state.chat_history[-1]["query"] != query:
                    st.session_state.chat_history.append(chat_entry)


    if st.session_state.show_cost_options:

        st.write("You have selected a cost-saving strategy. Please choose a specific approach to proceed further.")

        # Main cost-saving strategy selection using radio buttons
        cost_strategy = st.radio(
            "Select a cost-saving strategy:",
            options=["Cost Consolidation", "Network Optimization", "Min Order Rule"],
            key="cost_strategy",
            index = None
        )

        # Handle cost consolidation selection
        if cost_strategy == "Cost Consolidation":
            st.session_state.cost_consolidation_type = "Cost Consolidation"

            st.write("You have selected 'Cost Consolidation.' Please choose a specific approach to continue.")

            # Sub-options for Cost Consolidation using radio buttons
            consolidation_approach = st.radio(
                "Select an approach:",
                options=["Static", "Dynamic"],
                key="consolidation_approach",
                index = None
            )



            if consolidation_approach == "Static":
                st.session_state.static_clicked = True

                @st.cache_data
                def load_data():
                    df = pd.read_excel('Complete Input.xlsx', sheet_name='Sheet1')
                    rate_card = pd.read_excel('Cost per pallet.xlsx')
                    df['SHIPPED_DATE'] = pd.to_datetime(df['SHIPPED_DATE'], dayfirst=True)
                    rate_card_ambient = pd.read_excel('Complete Input.xlsx', sheet_name='AMBIENT')
                    rate_card_ambcontrol = pd.read_excel('Complete Input.xlsx', sheet_name='AMBCONTROL')
                    return df, rate_card, rate_card_ambient, rate_card_ambcontrol

                    # Rename df to avoid conflicts

                data_df, rate_card, rate_card_ambient, rate_card_ambcontrol = load_data()

                st.session_state.dynamic_clicked = False
                st.write("You have selected the **Static** way of approaching the query. Here's the suitable result.")
                st.write("Choose number of delivery day scenarios")

                delivery_options = st.radio(
                    "Select an approach:",
                    options=["5 days delivery scenario", "4 days delivery scenario", "3 days delivery scenario",
                             "2 days delivery scenario",
                             "1 day delivery scenario"],
                    key="delivery_options",
                    index=None
                )

                # Update session state with the selected delivery option
                if delivery_options:
                    st.session_state.selected_delivery_option = delivery_options

                # Display the selected delivery option
                if st.session_state.selected_delivery_option:
                    st.write(f"You have selected {st.session_state.selected_delivery_option}")

                    # Define the delivery scenarios
                    scenarios = {
                        "5 days delivery scenario": [
                            "Mon_Tue_Wed_Thu_Fri", "Mon_Tue_Wed_Thu_Sat", "Mon_Tue_Wed_Thu_Sun",
                            "Mon_Tue_Wed_Fri_Sat", "Mon_Tue_Wed_Fri_Sun", "Mon_Tue_Thu_Fri_Sat",
                            "Mon_Tue_Thu_Fri_Sun", "Mon_Wed_Thu_Fri_Sat", "Mon_Wed_Thu_Fri_Sun",
                            "Tue_Wed_Thu_Fri_Sat", "Tue_Wed_Thu_Fri_Sun"
                        ],
                        "4 days delivery scenario": [
                            "Mon_Tue_Wed_Thu", "Mon_Tue_Wed_Fri", "Mon_Tue_Wed_Sat", "Mon_Tue_Wed_Sun",
                            "Mon_Tue_Thu_Fri", "Mon_Tue_Thu_Sat", "Mon_Tue_Thu_Sun", "Mon_Wed_Thu_Fri",
                            "Mon_Wed_Thu_Sat", "Mon_Wed_Thu_Sun", "Tue_Wed_Thu_Fri", "Tue_Wed_Thu_Sat",
                            "Tue_Wed_Thu_Sun"
                        ],
                        "3 days delivery scenario": [
                            "Mon_Tue_Wed", "Mon_Tue_Thu", "Mon_Tue_Fri", "Mon_Tue_Sat", "Mon_Tue_Sun",
                            "Mon_Wed_Thu", "Mon_Wed_Fri", "Mon_Wed_Sat", "Mon_Wed_Sun",
                            "Tue_Wed_Thu", "Tue_Wed_Fri", "Tue_Wed_Sat", "Tue_Wed_Sun"
                        ],
                        "2 days delivery scenario": [
                            "Mon_Tue", "Mon_Wed", "Mon_Thu", "Mon_Fri", "Mon_Sat", "Mon_Sun",
                            "Tue_Wed", "Tue_Thu", "Tue_Fri", "Tue_Sat", "Tue_Sun",
                            "Wed_Thu", "Wed_Fri", "Wed_Sat", "Wed_Sun"
                        ],
                        "1 day delivery scenario": [
                            "Only_Mon", "Only_Tue", "Only_Wed", "Only_Thu", "Only_Fri", "Only_Sat", "Only_Sun"
                        ]
                    }

                    # Display the scenarios based on the selected delivery option
                    selected_scenarios = scenarios.get(st.session_state.selected_delivery_option, [])

                    if selected_scenarios:
                        st.write(
                            f"The possible {st.session_state.selected_delivery_option.split()[0]}-day scenarios are : ")
                        scenarios_single_line = " ,  &nbsp;&nbsp;&nbsp;".join(selected_scenarios)
                        st.write(scenarios_single_line)
                        st.write(" ")

                        if 'parameters' not in st.session_state:
                            st.session_state.parameters = parameter_values.get_parameters_values(api_key, query)

                        static.find_cost_savings(data_df, rate_card, selected_scenarios ,parameters=st.session_state.parameters)


            elif consolidation_approach == "Dynamic":
                st.session_state.dynamic_clicked = True
                st.write("You have selected the **Dynamic** way of approaching the query.")

                file_path = 'Complete Input.xlsx' # Change this to your actual file path
                pallet_file ='Cost per pallet.xlsx'
                df = pd.read_excel(file_path)
                dff = pd.read_excel(pallet_file)
                columns_to_remove = ["SHIPPED_DATE", "Total Pallets" ,"Distance"]  # Replace with actual column names
                df_filtered = df.drop(columns=columns_to_remove)
                df_filtered["ORDER_ID"] = df_filtered["ORDER_ID"].astype(str)
                # Display in Streamlit
                with st.expander("📂  Other files used ", expanded=False):
                    st.dataframe(df_filtered.head())
                    st.write("~ **Cost per each pallet for different postcodes** ⬇️")
                    st.dataframe(dff.head())

                col1, col2, col3 = st.columns(3)

                with col1:
                    st.session_state.total_shipment_capacity = st.slider("TOTAL SHIPMENT CAPACITY", 26, 52, 46)

                with col3:
                    st.session_state.shipment_window_range = st.slider("SHIPMENT WINDOW", 0, 30, (2, 7))

                @st.cache_data
                def load_data():
                    df = pd.read_excel('Complete Input.xlsx', sheet_name='Sheet1')
                    df['SHIPPED_DATE'] = pd.to_datetime(df['SHIPPED_DATE'], dayfirst=True)
                    rate_card_ambient = pd.read_excel('Complete Input.xlsx', sheet_name='AMBIENT')
                    rate_card_ambcontrol = pd.read_excel('Complete Input.xlsx', sheet_name='AMBCONTROL')
                    return df, rate_card_ambient, rate_card_ambcontrol

                # Rename df to avoid conflicts
                data_df, rate_card_ambient, rate_card_ambcontrol = load_data()

                if 'parameters' not in st.session_state:
                    st.session_state.parameters = parameter_values.get_parameters_values(api_key,query)

                # Initialize session state for selected postcodes and customers if they don't exist
                if 'selected_postcodes' not in st.session_state:
                    st.session_state.selected_postcodes = st.session_state.parameters.get('selected_postcodes', [])
                if 'selected_customers' not in st.session_state:
                    st.session_state.selected_customers = st.session_state.parameters.get('selected_customers', [])

                if not st.session_state.selected_postcodes and st.session_state.selected_customers:
                    group_method = 'Customer Level'
                else:
                    group_method = st.radio("Consolidation Level", ('Post Code Level', 'Customer Level'))

                group_field = 'SHORT_POSTCODE' if group_method == 'Post Code Level' else 'NAME'

                if group_method == 'Post Code Level':
                    all_postcodes = st.checkbox("All Post Codes", value=False)

                    if not all_postcodes:
                        postcode_counts = data_df['SHORT_POSTCODE'].value_counts()
                        postcode_options = postcode_counts.index.tolist()
                        selected_postcodes = st.multiselect(
                            "Select Post Codes",
                            options=postcode_options,
                            default=st.session_state.selected_postcodes,
                            format_func=lambda x: f"{x} ({postcode_counts[x]})")

                        st.session_state.selected_postcodes=selected_postcodes

                else:  # Customer Level
                    all_customers = st.checkbox("All Customers", value=False)

                    if not all_customers:
                        customer_counts = data_df['NAME'].value_counts()
                        customer_options = customer_counts.index.tolist()
                        selected_customers = st.multiselect(
                            "Select Customers",
                            options=customer_options,
                            default=st.session_state.selected_customers,
                            format_func=lambda x: f"{x} ({customer_counts[x]})"
                        )
                        st.session_state.selected_customers = selected_customers

                # Filter the dataframe based on the selection
                if group_method == 'Post Code Level' and not all_postcodes:
                    if st.session_state.selected_postcodes:  # Only filter if some postcodes are selected
                        data_df = data_df[data_df['SHORT_POSTCODE'].isin(st.session_state.selected_postcodes)]

                elif group_method == 'Customer Level' and not all_customers:
                    if st.session_state.selected_customers:  # Only filter if some customers are selected
                        data_df = data_df[data_df['NAME'].isin(st.session_state.selected_customers)]

                # Handle dynamic execution
                if st.button("Run Simulation"):
                    if st.session_state.dynamic_clicked:
                        with st.spinner("Running cost optimization simulation. Please wait..."):
                            #parameters = parameter_values.get_parameters_values(api_key, query)

                            st.session_state.results = cost_cosnsolidation.run_cost_optimization_simulation(st.session_state.parameters,
                                                                                                            api_key , total_capacity =st.session_state.total_shipment_capacity,shipment_window = st.session_state.shipment_window_range , customers = st.session_state.selected_customers ,postcodes = st.session_state.selected_postcodes )
                            cost_cosnsolidation.cost_calculation(st.session_state.parameters, st.session_state.results['params'] ,total_capacity =st.session_state.total_shipment_capacity)



        elif cost_strategy == "Network Optimization":
            st.session_state.cost_consolidation_type = "Network Optimization"
            st.write("You have selected 'Network Optimization.' Here's the approach:")

        elif cost_strategy == "Min Order Rule":
            st.session_state.cost_consolidation_type = "Min Order Rule"
            st.write("You have selected 'Min Order Rule.' Here's the approach:")

    # Display analysis history with download and delete options
    if st.session_state.chat_history:
        col1, col2 = st.columns([6, 2])
        with col1:
            st.subheader("📜 Analysis History")
        with col2:
            download_word_doc()

        # Iterate through history in reverse order with index tracking
        for idx, chat in enumerate(reversed(st.session_state.chat_history)):
            # Calculate the actual index in the original list
            original_idx = len(st.session_state.chat_history) - idx - 1

            # Create two columns for each analysis entry
            hist_col1, hist_col2 = st.columns([20, 1])

            with hist_col1:
                with st.expander(
                        f"Query {len(st.session_state.chat_history) - idx}: {chat['query'][:50]}...",
                        expanded=False
                ):
                    st.markdown("**🔍 Query:**")
                    st.write(chat['query'])

                    st.markdown(f"**📊 Data Source:** {chat.get('data_source', 'Not specified')}")

                    if chat['approach']:
                        st.markdown("**🎯 Approach:**")
                        st.write(chat['approach'])

                    if chat['answer']:
                        st.markdown("**💡 Results:**")
                        st.write(chat['answer'])

                    if chat['figure']:
                        st.pyplot(chat['figure'])

            with hist_col2:
                # Add delete button for each entry
                if st.button("🗑️", key=f"delete_{original_idx}"):
                    st.session_state.chat_history.pop(original_idx)
                    st.rerun()  # Rerun the app to refresh the display


if __name__ == "__main__":
    main()